from typing import Any, Tuple, Dict, List
from enum import Enum
import traceback
import erniebot
import unicodedata
import gradio as gr
from docx import Document
import requests
from docx.shared import Inches
import os

access_token_global = None
ak_global = None
sk_global = None

def query(message: str, prompt: str = None) -> Tuple[str, Any]:
    model_name = "ernie-4.0"  # 替换为实际使用的模型名称
    messages = [{"role": "user", "content": message}]
    functions = None  # 如果有特定的函数调用需求，可以在此处定义

    response = erniebot.ChatCompletion.create(
        _config_=dict(
            api_type="aistudio",
            access_token=access_token_global,
            ),
        model=model_name,
        messages=messages,
        functions=functions,
    )

    # 假设 response 返回的内容中包含我们需要的回复文本
    reply = response.get_result()
    return unicodedata.normalize('NFKC', reply), reply

def generate_image(description: str) -> str:
    response = erniebot.Image.create(
        _config_=dict(
            api_type="yinian",
            ak = ak_global,
            sk = sk_global,
            ),
        model="ernie-vilg-v2", 
        prompt=description, 
        width=1280, 
        height=720, 
        version="v2", 
        image_num=1
    )
    img_url = response.data["sub_task_result_list"][0]["final_image_list"][0]["img_url"]
    #print(img_url)
    return img_url

class STATE(Enum):
    init = 0
    gen_outline = 1
    gen_detail = 2
    rating = 3
    polish = 4
    gen_images = 5
    done = 100
    failed = -1

# theme
prompt_thesis_query = '''如果我想要写一篇关于{0}的论文，至少包含以下几个方面的内容：摘要、引言、方法、结论。\
请用简单的语言描述一下每个部分的内容应该怎么写，并根据主题适当添加例如相关工作、实验过程、参考文献等必要板块。\
你无需回答其它任何内容，也无需考虑可选部分，但需要遵循如下格式：\
请用"\n\n"分隔每个板块，例如：“摘要\n\n这是摘要的内容\n\n引言\n\n这是引言的内容\n\n...”\
'''
# theme section requirement extra
prompt_detail_query = '''如果我想要写一篇关于{0}的论文，请为我生成这篇论文的{1}部分，要求如下：{2}。\
你无需回答其它任何内容，但需要包含两个部分：\
第一个部分是用几句话简单的说明{1}的内容，第二个部分是{1}的具体内容，第二个部分尽量详细，字数越多越好。\
两个部分当中用"\n\n"分隔，且不需要任何提示语。\
'''
# theme thesis_content
prompt_rating_query = '''这是我写的关于{0}的论文\n\n{1}\n\n\
如果从专业性、完整性等角度评价，你会给这篇论文打多少分？对论文提出修改意见，你会怎么评价？\
你无需回答其它任何内容，但需要遵循如下格式：\
请用“\n\n”分隔每个板块，每个板块给出评分（1至5分）和原因，例如：\
“专业性：4分\n\n对专业性的修改意见\n\n完整性：2分\n\n对完整性的修改意见\n\n总体：3分\n\n对整篇论文的修改意见”\
'''
# theme thesis_simple section rating_content section_content
prompt_detail_finetune = '''这是我写的关于{0}的论文大纲\n\n{1}\n\n\
专家评价这篇论文并提出了一些修改意见，例如{3}。你可以帮我修改一下这篇论文的{2}部分吗？我是这样写的：{4}。\
你无需回答其它任何内容，但需要包含两个部分：\
第一个部分是用几句话简单的说明{2}的内容，第二个部分是{2}的具体内容。\
两个部分当中用"\n\n"分隔。\
'''

# prompt for generating images
prompt_image_query = '''请从以下内容中找出需要图片进行佐证的地方，并为这些地方生成配图（若没有需要配图的地方则不需要配图，尽量少配图）：\n\n{0}\n\n\
你无需回答其它任何内容，但需要遵循如下格式：\
如果需要配图，请仅回答一个详细的图片描述和一个图片题目，例如：“这是一个示例图片描述\n\n示例图片题目”。\
如果不需要配图，请仅回答“无需要配图”这5个字。\
每个部分间请用“\n\n”分隔。\
'''

def gen_outline(theme: str) -> Tuple[STATE, Dict[str, str], Dict[str, str]]:
    x = query(prompt_thesis_query.format(theme))
    results = x[0].strip().split('\n\n')
    thesis_req = {}
    thesis = {'主题': theme}
    while 1:
        try:
            key, value = results.pop(0).strip(), results.pop(0).strip()
            thesis_req[key] = value
        except:
            pass
        if not results:
            break
    try:
        assert '摘要' in thesis_req.keys(), '摘要'
        #assert '方法' in thesis_req.keys(), '方法'
        #assert '结论' in thesis_req.keys(), '结论'
    except:
        return STATE.init, {}, {}
    return STATE.gen_detail, thesis_req, thesis

def gen_detail(theme: str, thesis_req: Dict[str, str]) -> Tuple[STATE, Dict[str, str], Dict[str, str]]:
    thesis = {'主题': theme}
    thesis_simple = {}
    state = STATE.gen_detail
    for key in thesis_req.keys():
        print(f'{state.name} {key}')
        try:
            x = query(prompt_detail_query.format(theme, key, thesis_req[key]))
            results = x[0].strip().split('\n\n')
            thesis_simple[key] = results[0]
            thesis[key] = " ".join(results[1:])
        except:
            pass
    return STATE.rating, thesis, thesis_simple

def gen_rating(theme: str, thesis: Dict[str, str]) -> Tuple[STATE, Dict[str, str]]:
    x = query(prompt_rating_query.format(theme, ''.join([f'{k}：{v}' for k, v in thesis.items()])))
    results = x[0].strip().split('\n\n')
    rating = {}
    while 1:
        try:
            key, value = results.pop(0).strip(), results.pop(0).strip()
            rating[key] = value
        except:
            print(f'Ignored: {results}')
        if not results:
            break
    return STATE.polish, rating

def polish(
    theme: str, 
    thesis: Dict[str, str], 
    thesis_simple: Dict[str, str], 
    rating: Dict[str, str]
) -> Tuple[STATE, Dict[str, str], Dict[str, str]]:
    for key, value in thesis.items():
        try:
            x = query(prompt_detail_finetune.format(
                theme, 
                ''.join([f'{k}：{v}' for k, v in thesis_simple.items()]),
                key,
                ''.join([f'{k}：{v}' for k, v in rating.items()]),
                value
            ))
            results = x[0].strip().split('\n\n')
            thesis_simple[key] = results[0]
            thesis[key] = " ".join(results[1:])
        except:
            print(f'Ignored: {key} {value}')
    return STATE.gen_images, thesis, thesis_simple

def gen_images(thesis: Dict[str, str]) -> Tuple[STATE, Dict[str, str], List[Tuple[str, str, str]]]:
    images = []
    for key, value in thesis.items():
        if key == '主题' or key == '引言' or key == '摘要' or key == '参考文献':
            continue
        try:
            x = query(prompt_image_query.format(value))
            #print(x)
            results = x[0].strip().split('\n\n')
            image_index = 1
            i = 0
            while i < len(results) - 1:
                if results[i].strip() == "无需要配图":
                    i += 1
                    continue
                description = results[i].strip()
                title = results[i + 1].strip()
                #print('title done')
                image_url = generate_image(description)
                #print('url done')
                images.append((key, title, image_url))
                #thesis[key] += f'\n\n[插图：{title}]\n'
                image_index += 1
                i += 2
        except:
            continue
    return STATE.done, thesis, images

def main_loop(theme):
    state = STATE.init
    thesis = {}
    thesis_simple = {}
    thesis_req = {}
    rating = {}
    images = []
    while state != STATE.done and state != STATE.failed:
        print(state.name)
        try:
            if state == STATE.init:
                state = STATE.gen_outline
            elif state == STATE.gen_outline:
                state, thesis_req, thesis = gen_outline(theme)
            elif state == STATE.gen_detail:
                state, thesis, thesis_simple = gen_detail(theme, thesis_req)
            elif state == STATE.rating:
                state, rating = gen_rating(theme, thesis)
            elif state == STATE.polish:
                state, thesis, thesis_simple = polish(theme, thesis, thesis_simple, rating)
            elif state == STATE.gen_images:
                state, thesis, images = gen_images(thesis)

        except Exception as e:
            state = STATE.failed
            print(traceback.format_exc())

    return thesis, images

def download_image(url, filename):
    try:
        '''response = requests.get(url, stream=True)
        response.raise_for_status()
        with open(filename, 'wb') as file:
            for chunk in response.iter_content(1024):
                file.write(chunk)'''
        response = requests.get(url)
        with open(filename, "wb") as f:
            f.write(response.content)
    except Exception as e:
        print(f"Error occurred when downloading image from {url}, error message: {e}")

def create_word_document(thesis, images, filename="thesis.docx"):
    doc = Document()
    image_folder = "./images/"
    # 检查并创建图片保存的文件夹
    if not os.path.exists(image_folder):
        os.makedirs(image_folder)
    for key, value in thesis.items():
        doc.add_heading(key, level=1)
        paragraphs = value.split("\n\n")
        for paragraph in paragraphs:
            doc.add_paragraph(paragraph)
            for (image_key, title, url) in images:
                if image_key != key:
                    continue
                else:
                    image_path = os.path.join(image_folder, f"{title}.jpg")
                    download_image(url, image_path)
                    #print(image_key, title, url)
                    doc.add_picture(image_path, width=Inches(5.0))
                    doc.add_paragraph(title)
    doc.save(filename)

def create_ui_and_launch():
    with gr.Blocks(title="论文生成器", theme=gr.themes.Soft()) as blocks:
        gr.Markdown("# 基于百度erniebot的论文生成器")
        
        def _infer(
            theme, access_key, secret_key, access_token
        ):
            global access_token_global, ak_global, sk_global
            access_key = access_key.strip()
            secret_key = secret_key.strip()
            access_token = access_token.strip()
            history = []
            if (access_key == "" or secret_key == "") and access_token == "":
                raise gr.Error("需要填写正确的AK/SK或access token，不能为空")
            if theme.strip() == "":
                raise gr.Error("输入不能为空，请在清空后重试")
            
            access_token_global = access_token
            ak_global = access_key
            sk_global = secret_key

            state = STATE.init
            thesis = {}
            thesis_simple = {}
            thesis_req = {}
            rating = {}
            images = []
            while state != STATE.done and state != STATE.failed:
                print(state.name)
                str_state = str(state.name)
                feedback = 'null'
                try:
                    if state == STATE.init:
                        state = STATE.gen_outline
                    elif state == STATE.gen_outline:
                        state, thesis_req, thesis = gen_outline(theme)
                        feedback = str(thesis_req)
                    elif state == STATE.gen_detail:
                        state, thesis, thesis_simple = gen_detail(theme, thesis_req)
                        feedback = str(thesis)
                    elif state == STATE.rating:
                        state, rating = gen_rating(theme, thesis)
                        feedback = str(rating)
                    elif state == STATE.polish:
                        state, thesis, thesis_simple = polish(theme, thesis, thesis_simple, rating)
                        feedback = str(thesis)
                    elif state == STATE.gen_images:
                        state, thesis, images = gen_images(thesis)
                        feedback = str(images)
                except Exception as e:
                    state = STATE.failed
                    print(traceback.format_exc())
                history.append((str_state, feedback))
                str_state_now = str(state.name)
                yield str_state_now, history, None
            if state == STATE.done:
                create_word_document(thesis, images, filename="thesis.docx")
                yield str_state_now, history, "thesis.docx"
            else:
                yield str_state_now, history, None
            
        with gr.Row():
            with gr.Column(scale=1):
                access_key = gr.Textbox(
                    label="AK", info="yinian API的API Key", type="password"
                )
                secret_key = gr.Textbox(
                    label="SK", info="yinian API的Secret Key", type="password"
                )
                access_token = gr.Textbox(
                    label="Access Token", info="AI Studio的access token", type="password"
                )
                theme = gr.Textbox(label="论文主题", placeholder="请输入论文主题...")
            with gr.Column(scale=4):
                progress_display = gr.Textbox(label="进度")
                output_text = gr.Chatbot(label="生成内容")
                generate_btn = gr.Button("生成论文")
                download_btn = gr.File(label="下载论文", interactive=False)

        generate_btn.click(
            _infer,
            inputs=[
                theme,
                access_key,
                secret_key,
                access_token,
            ],
            outputs=[
                progress_display,
                output_text,
                download_btn,
            ],
        )
    blocks.launch()

if __name__ == '__main__':
    '''thesis, images = main_loop('自然语言处理')
    print(thesis)
    print(images)'''
    create_ui_and_launch()